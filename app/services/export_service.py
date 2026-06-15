"""Export Service for PDF, Excel, and Word exports"""
import os
from datetime import datetime
from io import BytesIO
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from database.models import Quotation, Invoice

class ExportService:
    """Service for exporting quotations and invoices"""
    
    @staticmethod
    def export_quotation_pdf(quotation_id, output_path=None):
        """
        Export quotation as PDF
        
        Args:
            quotation_id: ID of quotation to export
            output_path: Path to save PDF (optional)
        
        Returns:
            PDF file bytes
        """
        quotation = Quotation.query.get(quotation_id)
        if not quotation:
            raise ValueError(f"Quotation {quotation_id} not found")
        
        # Create PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        styles = getSampleStyleSheet()
        
        # Title
        title = Paragraph(f"<b>Quotation #{quotation.quotation_number}</b>", styles['Heading1'])
        elements.append(title)
        elements.append(Spacer(1, 0.3*inch))
        
        # Customer Info
        customer_info = f"""
        <b>Customer:</b> {quotation.customer.name}<br/>
        <b>Email:</b> {quotation.customer.email}<br/>
        <b>Company:</b> {quotation.customer.company}<br/>
        <b>Phone:</b> {quotation.customer.phone}<br/>
        <b>Date:</b> {quotation.created_at.strftime('%Y-%m-%d')}
        """
        elements.append(Paragraph(customer_info, styles['Normal']))
        elements.append(Spacer(1, 0.3*inch))
        
        # Items Table
        items_data = [['Material', 'Quantity', 'Unit', 'Unit Price', 'Total']]
        for item in quotation.items:
            items_data.append([
                item.material.name,
                f"{item.quantity}",
                item.unit,
                f"${item.unit_price:.2f}",
                f"${item.line_total:.2f}"
            ])
        
        # Add totals row
        items_data.append([
            '',
            '',
            '',
            'Subtotal:',
            f"${quotation.subtotal:.2f}"
        ])
        items_data.append([
            '',
            '',
            '',
            'Tax:',
            f"${quotation.tax_amount:.2f}"
        ])
        items_data.append([
            '',
            '',
            '',
            'Total:',
            f"${quotation.total_amount:.2f}"
        ])
        
        table = Table(items_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        elements.append(table)
        elements.append(Spacer(1, 0.3*inch))
        
        # Notes
        if quotation.notes:
            elements.append(Paragraph(f"<b>Notes:</b> {quotation.notes}", styles['Normal']))
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        
        # Save to file if path provided
        if output_path:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            with open(output_path, 'wb') as f:
                f.write(buffer.getvalue())
        
        return buffer.getvalue()
    
    @staticmethod
    def export_quotation_excel(quotation_id, output_path=None):
        """
        Export quotation as Excel
        
        Args:
            quotation_id: ID of quotation to export
            output_path: Path to save Excel file (optional)
        
        Returns:
            Excel file bytes
        """
        quotation = Quotation.query.get(quotation_id)
        if not quotation:
            raise ValueError(f"Quotation {quotation_id} not found")
        
        # Create workbook
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Quotation"
        
        # Set column widths
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 12
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['E'].width = 15
        
        # Header
        ws['A1'] = f"Quotation #{quotation.quotation_number}"
        ws['A1'].font = Font(bold=True, size=14)
        
        # Customer Info
        ws['A3'] = "Customer Information"
        ws['A3'].font = Font(bold=True)
        ws['A4'] = f"Name: {quotation.customer.name}"
        ws['A5'] = f"Email: {quotation.customer.email}"
        ws['A6'] = f"Company: {quotation.customer.company}"
        ws['A7'] = f"Date: {quotation.created_at.strftime('%Y-%m-%d')}"
        
        # Items
        ws['A9'] = "Item Details"
        ws['A9'].font = Font(bold=True)
        
        headers = ['Material', 'Quantity', 'Unit', 'Unit Price', 'Total']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=10, column=col)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
        
        row = 11
        for item in quotation.items:
            ws.cell(row=row, column=1).value = item.material.name
            ws.cell(row=row, column=2).value = item.quantity
            ws.cell(row=row, column=3).value = item.unit
            ws.cell(row=row, column=4).value = item.unit_price
            ws.cell(row=row, column=5).value = item.line_total
            row += 1
        
        # Totals
        row += 1
        ws.cell(row=row, column=4).value = "Subtotal:"
        ws.cell(row=row, column=4).font = Font(bold=True)
        ws.cell(row=row, column=5).value = quotation.subtotal
        
        row += 1
        ws.cell(row=row, column=4).value = "Tax:"
        ws.cell(row=row, column=4).font = Font(bold=True)
        ws.cell(row=row, column=5).value = quotation.tax_amount
        
        row += 1
        ws.cell(row=row, column=4).value = "Total:"
        ws.cell(row=row, column=4).font = Font(bold=True)
        ws.cell(row=row, column=5).value = quotation.total_amount
        
        # Save to file
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        if output_path:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            wb.save(output_path)
        
        return buffer.getvalue()
    
    @staticmethod
    def export_quotation_word(quotation_id, output_path=None):
        """
        Export quotation as Word document
        
        Args:
            quotation_id: ID of quotation to export
            output_path: Path to save Word document (optional)
        
        Returns:
            Word document file bytes
        """
        quotation = Quotation.query.get(quotation_id)
        if not quotation:
            raise ValueError(f"Quotation {quotation_id} not found")
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f"Quotation #{quotation.quotation_number}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Customer info
        doc.add_heading("Customer Information", level=2)
        doc.add_paragraph(f"Name: {quotation.customer.name}")
        doc.add_paragraph(f"Email: {quotation.customer.email}")
        doc.add_paragraph(f"Company: {quotation.customer.company}")
        doc.add_paragraph(f"Phone: {quotation.customer.phone}")
        doc.add_paragraph(f"Date: {quotation.created_at.strftime('%Y-%m-%d')}")
        
        # Items
        doc.add_heading("Items", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Material'
        hdr_cells[1].text = 'Quantity'
        hdr_cells[2].text = 'Unit'
        hdr_cells[3].text = 'Unit Price'
        hdr_cells[4].text = 'Total'
        
        for item in quotation.items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.material.name
            row_cells[1].text = str(item.quantity)
            row_cells[2].text = item.unit
            row_cells[3].text = f"${item.unit_price:.2f}"
            row_cells[4].text = f"${item.line_total:.2f}"
        
        # Totals
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(f"Subtotal: ${quotation.subtotal:.2f}")
        doc.add_paragraph(f"Tax ({quotation.tax_percentage}%): ${quotation.tax_amount:.2f}")
        doc.add_paragraph(f"Total: ${quotation.total_amount:.2f}").runs[0].font.bold = True
        
        # Notes
        if quotation.notes:
            doc.add_heading("Notes", level=2)
            doc.add_paragraph(quotation.notes)
        
        # Save to file
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        if output_path:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
        
        return buffer.getvalue()
